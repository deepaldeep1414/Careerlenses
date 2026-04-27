import os
import uuid
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.makedirs("outputs", exist_ok=True)

def generate_docx(data: dict) -> str:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    resume = data.get("structured_resume", {})
    enhanced = data.get("enhanced", {})

    name = resume.get("name") or data.get("name", "Your Name")
    email = resume.get("email") or data.get("email", "")
    phone = resume.get("phone") or data.get("phone", "")
    linkedin = resume.get("linkedin", "")
    summary = enhanced.get("summary") or resume.get("summary", "")
    skills = enhanced.get("enhanced_skills") or resume.get("skills", [])
    experience = enhanced.get("enhanced_experience") or resume.get("experience", [])
    education = resume.get("education", [])
    projects = enhanced.get("enhanced_projects") or resume.get("projects", [])

    def add_section_heading(title):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(name)
    run.bold = True
    run.font.size = Pt(18)

    # Contact
    contact_parts = [x for x in [email, phone, linkedin] if x]
    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run("  |  ".join(contact_parts)).font.size = Pt(10)

    # Summary
    if summary:
        add_section_heading("Professional Summary")
        p = doc.add_paragraph(summary)
        p.runs[0].font.size = Pt(10)

    # Skills
    if skills:
        add_section_heading("Technical Skills")
        p = doc.add_paragraph(", ".join(skills))
        p.runs[0].font.size = Pt(10)

    # Experience
    if experience:
        add_section_heading("Experience")
        for job in experience:
            p = doc.add_paragraph()
            r1 = p.add_run(job.get("company", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            p.add_run(
                f"  |  {job.get('title', '')}  |  {job.get('dates', '')}"
            ).font.size = Pt(10)
            for bullet in job.get("bullets", []):
                b = doc.add_paragraph(style="List Bullet")
                b.add_run(bullet).font.size = Pt(10)

    # Projects
    if projects:
        add_section_heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            r1 = p.add_run(proj.get("name", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            desc = proj.get("description", "")
            if desc:
                p.add_run(f" — {desc}").font.size = Pt(10)
            for bullet in proj.get("bullets", []):
                b = doc.add_paragraph(style="List Bullet")
                b.add_run(bullet).font.size = Pt(10)

    # Education
    if education:
        add_section_heading("Education")
        for edu in education:
            p = doc.add_paragraph()
            r1 = p.add_run(edu.get("school", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            p.add_run(
                f"  |  {edu.get('degree', '')}  |  {edu.get('year', '')}"
            ).font.size = Pt(10)

    file_id = str(uuid.uuid4())
    output_path = f"outputs/{file_id}.docx"
    doc.save(output_path)
    return file_id